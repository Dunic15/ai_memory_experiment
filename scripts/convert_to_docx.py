#!/usr/bin/env python3
"""
Convert the markdown document to .docx format.
Requires python-docx library.
"""

import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: python-docx is not installed.")
    print("Please install it with: pip install python-docx")
    sys.exit(1)

def parse_markdown(markdown_file):
    """Parse markdown file line by line."""
    with open(markdown_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    return lines

def add_paragraph_with_formatting(doc, text, style='Normal', bold=False, italic=False, size=None):
    """Add a paragraph with formatting."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    return p

def create_docx(markdown_file, output_file):
    """Create a .docx file from markdown."""
    print(f"Reading markdown file: {markdown_file}")
    lines = parse_markdown(markdown_file)
    
    print(f"Creating .docx file: {output_file}")
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Main heading (##)
        if line.startswith('## '):
            title = line[3:].strip()
            doc.add_heading(title, level=1)
        
        # Subheading (###)
        elif line.startswith('### '):
            title = line[4:].strip()
            doc.add_heading(title, level=2)
        
        # Horizontal rule (---)
        elif line.strip() == '---':
            doc.add_paragraph()  # Add spacing
        
        # Bold text (questions)
        elif line.startswith('**') and line.endswith('**'):
            text = line.replace('**', '').strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
        
        # Options with checkmark (correct answer)
        elif re.match(r'^\s*✓\s*[a-d]\)', line):
            text = line.strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.color.rgb = RGBColor(0, 128, 0)  # Green
        
        # Regular options
        elif re.match(r'^\s*[a-d]\)', line):
            text = line.strip()
            doc.add_paragraph(text)
        
        # Empty line
        elif not line.strip():
            doc.add_paragraph()
        
        # Regular paragraph
        else:
            # Collect multi-line paragraphs
            para_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip() and not lines[i].startswith('#') and not lines[i].startswith('**') and not re.match(r'^\s*[a-d]\)', lines[i]) and lines[i].strip() != '---':
                para_lines.append(lines[i].rstrip())
                i += 1
            i -= 1  # Back up one line
            
            para_text = ' '.join(para_lines).strip()
            if para_text:
                doc.add_paragraph(para_text)
        
        i += 1
    
    doc.save(output_file)
    print(f"✓ Document saved to: {output_file}")

def main():
    """Main function."""
    base_dir = Path(__file__).parent.parent
    
    # Convert document WITH answers
    markdown_file_with_answers = base_dir / "Articles_English_Chinese.md"
    output_file_with_answers = base_dir / "Articles_English_Chinese.docx"
    
    if markdown_file_with_answers.exists():
        print(f"Converting document WITH answers...")
        create_docx(markdown_file_with_answers, output_file_with_answers)
        print(f"✓ Successfully created: {output_file_with_answers}")
    else:
        print(f"Warning: {markdown_file_with_answers} not found!")
    
    # Convert document WITHOUT answers
    markdown_file_no_answers = base_dir / "Articles_English_Chinese_NoAnswers.md"
    output_file_no_answers = base_dir / "Articles_English_Chinese_NoAnswers.docx"
    
    if markdown_file_no_answers.exists():
        print(f"\nConverting document WITHOUT answers...")
        create_docx(markdown_file_no_answers, output_file_no_answers)
        print(f"✓ Successfully created: {output_file_no_answers}")
    else:
        print(f"Warning: {markdown_file_no_answers} not found!")
    
    if not markdown_file_with_answers.exists() and not markdown_file_no_answers.exists():
        print(f"\nError: No markdown files found!")
        print("Please run generate_chinese_document.py first.")
        sys.exit(1)
    
    print(f"\n✓ All documents converted successfully!")
    print(f"  You can now copy these files to: Tesi/Cina tesi/Thesis China/")

if __name__ == "__main__":
    main()

